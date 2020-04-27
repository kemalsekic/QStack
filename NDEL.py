import docx
import os
import shutil
import shelve
import pprint
import subprocess
import webbrowser
import configStuff
import driverThings
import deviceDetail
from os import system, name
import logging
import loginPage
import stagingFile
import driverConfig

driver = driverConfig.Driver
doc = docx.Document()

def clearS():
    if name == 'nt': 
        _ = system('cls') 

#Create the Document and add headers
doc.add_heading('Cubby - Staging Device', 0)
doc.add_paragraph("Username: " + stagingFile.uName + "\t\t\tPassword: " + stagingFile.pWord)
doc.add_paragraph("Account #: " + stagingFile.accNum + "\t\t\t\tSSN: " + stagingFile.ssn)
doc.add_paragraph("Name: " + stagingFile.firstName + " " + stagingFile.lastName
+ "\t\t\t\t\tDOB: " + stagingFile.dobMonth + "/" + stagingFile.dobDay + "/" + stagingFile.dobYear)
doc.add_paragraph("Address: " + stagingFile.address + "\t\t\t\t\tAddress Line 2: " + stagingFile.address2)
doc.add_paragraph("City/State/Zip: " + stagingFile.city + "/" + stagingFile.state + "/" + stagingFile.zipCode
+ "\t\t\t\tMDN: " + stagingFile.mdn)
doc.add_paragraph("Invoice #: " + stagingFile.invoiceNum + "\t\t\t\t\t\tOrder #: " + stagingFile.orderNum)

def addScreenShot(imageName, headingText):
    configStuff.count += 1
    driver.save_screenshot(imageName + str(configStuff.count) + ".png")
    doc.add_heading("Step " + str(configStuff.count) + " - " + headingText, 2)
    doc.add_picture(imageName + str(configStuff.count) + '.png', width=docx.shared.Inches(6.5), height=docx.shared.Inches(3.66))
    os.remove('C:/Users/P2938214/Documents/Chemal/QStack/' + imageName + str(configStuff.count) + '.png')
    #source = 'C:/Users/P2938214/Documents/Chemal/QStack/' + (imageName + str(configStuff.count))
    #destination = 'C:/Users/P2938214/Documents/Chemal/QStack/screenshots'
    #shutil.move(source, destination)

clearS()
#docName = input("Name of document: ")
docName = "docStuff"

#Load site
driver.maximize_window()
driver.get('https://chtrgwy-uat.corp.chartercom.com/static/impersonation-portal/current/')

#Wait for element to display then click it
addScreenShot(docName, "Logging into site: https://chtrgwy-uat.corp.chartercom.com/static/impersonation-portal/current/")
#loginPage.loginToSite()

clearS()

#Finding elements on page
driverThings.NDEL_RetailWait()
addScreenShot(docName, "Page loaded")
driver.find_element_by_id("user-input-id").send_keys("retail")
driverThings.NDEL_RetailWaitButton()
driver.find_element_by_xpath("//*[@id='AP-CHTRGWY-Retail-SDP--card']/div[2]/div[2]/button[1]").click()

driver.get('https://chtrgwy-uat.corp.chartercom.com/static/gateway/beta/#/dashboard')
driverThings.waitForSearchButton()
addScreenShot(docName, "Clicking Search Button")
driver.find_element_by_xpath("//button[@class='MuiButtonBase-root MuiButton-root jss159 MuiButton-text']").click()

driverThings.waitForAccountNumButton()
driver.find_element_by_id("billerAccountNumber").send_keys(stagingFile.accNum)
addScreenShot(docName, "Entered Account Number")

driverThings.waitForGWFindAccButton()
driver.find_element_by_xpath("/html/body/div[6]/div[3]/div/div[2]/div/div/form/fieldset/div[3]/button[1]").click()

driverThings.wfverifyButton()
driver.find_element_by_xpath(driverThings.verifyButton).click()

driverThings.wfMobileButton()
driver.find_element_by_xpath(driverThings.mobileButton).click()

driverThings.wfTasksDD()
driver.find_element_by_xpath(driverThings.tasksDropDown).click()

driverThings.wfManageLinesSelect()
driver.find_element_by_xpath(driverThings.manageLinesSelect).click()

#driverThings.wfShopSelect()
#driver.find_element_by_xpath(driverThings.shopSelect).click()


"""
#Type into field
#8245100010088504@charter.net
#8245124000228185@charter.net
#8345780121546034

#Click the button to continue
driverThings.waitForSignInButton()
addScreenShot(docName, "Entered username and password")
signInButton = driver.find_element_by_xpath(".//*[contains(@class, 'button-lbg--primary')]")
signInButton.click()

driverThings.waitForDeviceButton()
driver.find_element_by_xpath(".//*[contains(@href, '/manage-account/devices')]").click()
addScreenShot(docName, "Select device")

driverThings.waitForAddALine()
addScreenShot(docName, "Adding a line")
addALineButton = driver.find_element_by_class_name('add-content')
driver.execute_script("arguments[0].click();", addALineButton)

driverThings.waitForDeviceList()
addScreenShot(docName, "Device list is displayed")
driver.find_element_by_xpath(".//*[(contains(@href, '/products/phone/iphone-11-pro-apple?colorName=Midnight%20Green'))]").click()

driverThings.waitForNextButton()
addScreenShot(docName, "Clicking next button here")
driver.find_element_by_xpath(".//*[contains(@class, 'button--primary next-btn')]").click()

driverThings.waitForSSNInput()
addScreenShot(docName, "SSN Input is loaded")

driver.find_element_by_id("ssn").send_keys(stagingFile.ssn)
addScreenShot(docName, "Entered SSN")

driver.find_element_by_id("month").send_keys(stagingFile.dobMonth)
driver.find_element_by_id("date").send_keys(stagingFile.dobDay)
driver.find_element_by_id("year").send_keys(stagingFile.dobYear)
addScreenShot(docName, "Entered date of birth: " + stagingFile.dobMonth + "/" + stagingFile.dobDay + "/" + stagingFile.dobYear)

driver.find_element_by_xpath("//span[text()='CONTINUE']").click()

driverThings.waitForByTheGigCard()
driver.find_element_by_xpath(".//h3[contains(@class, 'title')]").click()
addScreenShot(docName, "Selected By The Gig Plan")

#driverThings.waitForKeepMyNumber()
#driver.find_element_by_id("phone-number-Keepmyphonenumber").click()
#addScreenShot(docName, "")

driverThings.waitForDeviceNickName()
driver.find_element_by_xpath("//div[@class='personalize']//input").send_keys("HeyAll1")
addScreenShot(docName, "Entered Device Nick-Name")

#//p[@class='customizer-plan__price']
driverThings.waitForAddToCartButton()
driver.find_element_by_xpath("//div[@class='customizer-step']//button[contains(text(), 'Add to Cart')]").click()
addScreenShot(docName, "Added to Cart")

driverThings.waitForContinueToCartButton()
driver.find_element_by_xpath("//button[contains(text(), 'Continue to Cart')]").click()
addScreenShot(docName, "Clicked Continue to Cart")

driverThings.waitForSignOutButton()
driver.find_element_by_link_text('Sign Out').click()
driverThings.waitForAccessLink()
addScreenShot(docName, "Signed Out")
"""
doc.save(docName + '.docx')

clearS()
print("\n\n\t\tStaging Complete.\n\t\tCheck your document: [" + docName + ".docx]")
print("\n\n\t\tCubby\n\n\t\t\tBy QStack\n\n\n\n")
driver.quit()

